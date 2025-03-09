from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def formatar_legenda(documento, titulo, legenda):
    paragrafo = documento.add_paragraph()

    # Formatação do título
    run_titulo = paragrafo.add_run(titulo)
    run_titulo.bold = True
    run_titulo.font.size = Pt(12)
    run_titulo.font.color.rgb = RGBColor(0, 0, 0) # Black

    # Adicionando espaço entre título e legenda
    run_espaco = paragrafo.add_run(" ")

    # Formatação da legenda
    run_legenda = paragrafo.add_run(legenda)
    run_legenda.font.size = Pt(10)
    run_legenda.font.color.rgb = RGBColor(0, 0, 0) # Black

    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Exemplo de uso
documento = Document()
formatar_legenda(documento, "Figura 1. Exemplo de título", "Legenda da figura.")
documento.save("legenda_formatada.docx")
